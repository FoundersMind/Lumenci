import csv
import os
import shutil
from dataclasses import dataclass
from typing import List


@dataclass
class ParsedRow:
    row_index: int
    claim_text: str
    evidence_text: str
    reasoning_text: str
    strength: str = "weak"


class ParseError(Exception):
    pass


def _ext(path: str) -> str:
    return (os.path.splitext(path)[1] or "").lstrip(".").lower()


def _parse_csv(path: str) -> List[ParsedRow]:
    rows: List[ParsedRow] = []
    with open(path, "r", encoding="utf-8", newline="") as f:
        reader = csv.reader(f)
        for idx, r in enumerate(reader):
            if idx == 0 and len(r) >= 3 and "claim" in (r[0] or "").lower():
                continue
            if not r or all((c or "").strip() == "" for c in r):
                continue
            c0 = (r[0] if len(r) > 0 else "").strip()
            c1 = (r[1] if len(r) > 1 else "").strip()
            c2 = (r[2] if len(r) > 2 else "").strip()
            rows.append(
                ParsedRow(
                    row_index=len(rows) + 1,
                    claim_text=c0,
                    evidence_text=c1,
                    reasoning_text=c2,
                )
            )
    if not rows:
        raise ParseError("No rows detected in CSV.")
    return rows


def _parse_xlsx(path: str) -> List[ParsedRow]:
    try:
        from openpyxl import load_workbook
    except Exception as e:
        raise ParseError("Missing dependency: openpyxl") from e

    wb = load_workbook(path, data_only=True)
    ws = wb.active
    rows: List[ParsedRow] = []
    for ridx, row in enumerate(ws.iter_rows(values_only=True)):
        # try to skip header row
        if ridx == 0:
            v0 = str(row[0] or "").lower()
            if "claim" in v0 and len(row) >= 3:
                continue
        if not row:
            continue
        c0 = str(row[0] or "").strip()
        c1 = str(row[1] or "").strip() if len(row) > 1 else ""
        c2 = str(row[2] or "").strip() if len(row) > 2 else ""
        if not (c0 or c1 or c2):
            continue
        rows.append(
            ParsedRow(
                row_index=len(rows) + 1,
                claim_text=c0,
                evidence_text=c1,
                reasoning_text=c2,
            )
        )
    if not rows:
        raise ParseError("No rows detected in XLSX.")
    return rows


def _parse_docx(path: str) -> List[ParsedRow]:
    try:
        import docx
    except Exception as e:
        raise ParseError("Missing dependency: python-docx") from e

    d = docx.Document(path)
    if not d.tables:
        raise ParseError("No tables found in DOCX.")
    t = d.tables[0]
    rows: List[ParsedRow] = []
    for ridx, tr in enumerate(t.rows):
        cells = [c.text.strip() for c in tr.cells]
        if ridx == 0 and cells and "claim" in (cells[0] or "").lower():
            continue
        c0 = cells[0] if len(cells) > 0 else ""
        c1 = cells[1] if len(cells) > 1 else ""
        c2 = cells[2] if len(cells) > 2 else ""
        if not (c0 or c1 or c2):
            continue
        rows.append(
            ParsedRow(
                row_index=len(rows) + 1,
                claim_text=c0,
                evidence_text=c1,
                reasoning_text=c2,
            )
        )
    if not rows:
        raise ParseError("No rows detected in DOCX table.")
    return rows


def _tesseract_available() -> bool:
    return shutil.which("tesseract") is not None


def _ocr_image_to_text(path: str) -> str:
    try:
        from PIL import Image
    except Exception as e:
        raise ParseError("Missing dependency: pillow") from e
    try:
        import pytesseract
    except Exception as e:
        raise ParseError("Missing dependency: pytesseract") from e

    if not _tesseract_available():
        raise ParseError("Tesseract OCR is not installed or not on PATH.")

    img = Image.open(path)
    return pytesseract.image_to_string(img)


def _parse_image_heuristic(path: str) -> List[ParsedRow]:
    # Prototype OCR heuristic: extract full text and split into 3 chunks by simple markers.
    # Real-world claim charts vary widely; for demo we aim for a best-effort parse.
    text = _ocr_image_to_text(path)
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    if not lines:
        raise ParseError("No text detected in image.")

    # Heuristic: try to group lines into row blocks using blank-line gaps (already removed) and keywords.
    # We'll just create up to 10 rows by splitting every ~6 lines.
    chunk_size = 6
    rows: List[ParsedRow] = []
    for i in range(0, min(len(lines), 60), chunk_size):
        chunk = lines[i : i + chunk_size]
        if not chunk:
            continue
        c0 = chunk[0] if len(chunk) > 0 else ""
        c1 = " ".join(chunk[1:3]).strip() if len(chunk) > 1 else ""
        c2 = " ".join(chunk[3:]).strip() if len(chunk) > 3 else ""
        rows.append(
            ParsedRow(
                row_index=len(rows) + 1,
                claim_text=c0,
                evidence_text=c1,
                reasoning_text=c2,
            )
        )

    if not rows:
        raise ParseError("Unable to derive rows from OCR text.")
    return rows


def _parse_pdf(path: str) -> List[ParsedRow]:
    # Best-effort: try pdfplumber table extraction; else OCR first page image.
    try:
        import pdfplumber
    except Exception as e:
        raise ParseError("Missing dependency: pdfplumber") from e

    rows: List[ParsedRow] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages[:3]:
            try:
                tables = page.extract_tables() or []
            except Exception:
                tables = []
            for t in tables:
                for ridx, r in enumerate(t):
                    if ridx == 0 and r and "claim" in (str(r[0] or "").lower()):
                        continue
                    c0 = str(r[0] or "").strip() if len(r) > 0 else ""
                    c1 = str(r[1] or "").strip() if len(r) > 1 else ""
                    c2 = str(r[2] or "").strip() if len(r) > 2 else ""
                    if not (c0 or c1 or c2):
                        continue
                    rows.append(
                        ParsedRow(
                            row_index=len(rows) + 1,
                            claim_text=c0,
                            evidence_text=c1,
                            reasoning_text=c2,
                        )
                    )
            if rows:
                break
        if rows:
            return rows

        # OCR fallback (first page)
        first = pdf.pages[0] if pdf.pages else None
        if not first:
            raise ParseError("PDF is empty.")
        try:
            img = first.to_image(resolution=200).original
        except Exception as e:
            raise ParseError("Unable to rasterize PDF for OCR.") from e

    # OCR the rasterized image in-memory
    try:
        import pytesseract
    except Exception as e:
        raise ParseError("Missing dependency: pytesseract") from e
    try:
        from PIL import Image
    except Exception as e:
        raise ParseError("Missing dependency: pillow") from e

    if not _tesseract_available():
        raise ParseError("Tesseract OCR is not installed or not on PATH.")

    if not isinstance(img, Image.Image):
        img = Image.fromarray(img)
    text = pytesseract.image_to_string(img)
    if not text.strip():
        raise ParseError("No text detected in PDF via OCR.")
    # Reuse heuristic splitter
    tmp_path = None
    return _parse_image_heuristic(path)  # last-resort heuristic


def parse_claim_chart(path: str) -> List[ParsedRow]:
    ext = _ext(path)
    if ext in ("csv",):
        return _parse_csv(path)
    if ext in ("xlsx", "xlsm", "xltx", "xltm"):
        return _parse_xlsx(path)
    if ext in ("docx",):
        return _parse_docx(path)
    if ext in ("pdf",):
        return _parse_pdf(path)
    if ext in ("png", "webp", "jpg", "jpeg"):
        return _parse_image_heuristic(path)
    raise ParseError(f"Unsupported file type: .{ext}")


def extract_product_doc_text(path: str) -> str:
    """
    Best-effort plain-text extraction for product/technical documents.
    This is used as lightweight RAG context for the chat endpoint.
    """
    ext = _ext(path)
    if ext in ("txt",):
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            raise ParseError("Unable to read .txt document.") from e

    if ext in ("docx",):
        try:
            import docx
        except Exception as e:
            raise ParseError("Missing dependency: python-docx") from e
        d = docx.Document(path)
        parts: List[str] = []
        for p in d.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        # include table cell text too (common for spec sheets)
        for t in d.tables or []:
            for tr in t.rows:
                cells = [c.text.strip() for c in tr.cells if (c.text or "").strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    if ext in ("pdf",):
        try:
            import pdfplumber
        except Exception as e:
            raise ParseError("Missing dependency: pdfplumber") from e
        texts: List[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages[:8]:
                try:
                    t = (page.extract_text() or "").strip()
                except Exception:
                    t = ""
                if t:
                    texts.append(t)
        if texts:
            return "\n\n".join(texts)
        # OCR fallback
        return "\n".join([r.claim_text + "\n" + r.evidence_text + "\n" + r.reasoning_text for r in _parse_image_heuristic(path)])

    if ext in ("png", "webp", "jpg", "jpeg"):
        return _ocr_image_to_text(path)

    # For unknown types, do not pretend we can parse.
    raise ParseError(f"Unsupported doc type for extraction: .{ext}")

