import io
import json
import os
import re
import time
from typing import Any, Dict, List, Optional, Tuple

from django.conf import settings
from django.db import transaction
from django.db.models import Max, Prefetch
from django.http import FileResponse, JsonResponse
from django.shortcuts import get_object_or_404, render
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_GET, require_POST

from .models import Case, ChatMessage, ClaimChart, ClaimChartRow, ProductDoc, RowChange
from .parsing import ParseError, extract_product_doc_text, parse_claim_chart
from .strength_llm import sync_claim_chart_strengths, sync_one_row_strength
from docx import Document

try:
    from groq import Groq
except Exception:  # pragma: no cover
    Groq = None


SYSTEM_PROMPT = (
    "You are Lumenci Assistant, an AI specialized in patent infringement analysis. "
    "You help patent analysts refine claim charts by improving evidence quality and strengthening legal reasoning. "
    "The current claim chart is provided to you. When asked to improve a row, suggest specific, technically precise reasoning. "
    "When you suggest a change, clearly state: which row/element you're modifying, the old text, and your suggested new text. "
    "Format suggestions clearly so they can be accepted or rejected. If you cannot find strong evidence, say so and ask the analyst "
    "to add a technical/product document (file upload) or use **Add evidence from URL** to capture readable text from a public web page "
    "they are allowed to use. They may also paste excerpts into chat. "
    "When URL-captured text is present in the document snippets, treat it like other exhibit text. "
    "If reasoning should better address claim construction or element boundaries, say what is vague and propose clearer, legally precise wording.\n\n"
    "Keep responses concise and actionable.\n\n"
    "DOCUMENT ACCESS NOTE: The analyst may upload product/technical documents. You will be given extracted text snippets "
    "from those documents inside the system message. Do NOT claim you 'cannot read files' or 'cannot access uploads' if "
    "document text is provided. If a document has no extracted text, say it could not be processed yet and ask for a different format.\n\n"
    "Structured chart edits MUST appear inside <lumenci_suggestion_json> ... </lumenci_suggestion_json> exactly (not only as a naked JSON block). "
    "When the analyst says to apply changes, keep emitting that tagged JSON so the app can update the grid.\n\n"
    "IMPORTANT: At the end of your response, include a machine-readable JSON block between "
    "<lumenci_suggestion_json> and </lumenci_suggestion_json>. The JSON schema must be:\n"
    "{\n"
    '  "suggestions": [\n'
    "    {\n"
    '      "row_id": 1 | 2 | 3 | ...,\n'
    '      "field": "claim" | "evidence" | "reasoning",\n'
    '      "old_text": string,\n'
    '      "new_text": string\n'
    "    }\n"
    "  ],\n"
    '  "new_rows": [\n'
    "    {\n"
    '      "claim": string,\n'
    '      "evidence": string,\n'
    '      "reasoning": string\n'
    "    }\n"
    "  ]\n"
    "}\n"
    "Use \"new_rows\" ONLY when the analyst needs NEW claim chart rows that are not already present "
    "(e.g. missing features / additional claim elements). Each object must have at least one non-empty column. "
    "To refine text in an existing row, use \"suggestions\" with that row's row_id — do not duplicate the same element "
    "as a new row. If there are no new rows, return \"new_rows\": []. "
    "If you have no edits, return {\"suggestions\": [], \"new_rows\": []}.\n"
)

GROQ_RATE_LIMIT_USER_MESSAGE = (
    "Groq **rate limit** hit (requests or tokens per minute on your plan). "
    "Wait about a minute and send again, or use a lighter model in `.env`: "
    "`GROQ_MODEL=llama-3.1-8b-instant` or `meta-llama/llama-4-scout-17b-16e-instruct`. "
    "Limits: https://console.groq.com/docs/rate-limits"
)


def _groq_retry_after_seconds(exc: BaseException) -> int:
    try:
        r = getattr(exc, "response", None)
        if r is not None:
            h = r.headers.get("retry-after") if r.headers else None
            if h is not None:
                return min(max(int(float(h)), 1), 120)
    except Exception:
        pass
    return 6


STRENGTH_POLICY = (
    "\n\nROW STRENGTH (analyst judgement — treat as authoritative):\n"
    'Each row in the chart JSON has \"strength\": \"strong\" | \"weak\" | \"missing\".\n'
    "- Do NOT output JSON \"suggestions\" that modify any cell on a row where strength is \"strong\", "
    "unless the analyst clearly asked for optional alternates or a broad review (e.g. alternate wording, "
    "parallel argument, another angle, reframe, optional phrasing, review/all rows, entire chart, rewrite all).\n"
    "- For “fix weak reasoning”, “strengthen evidence”, or similar repair requests: ONLY suggest edits for rows "
    "with strength \"weak\" or \"missing\". If all relevant rows are already \"strong\", say so, congratulate brevity, "
    'return {\"suggestions\": [], \"new_rows\": []}, and do not fabricate weaknesses.\n'
    "- Do not rewrite strong rows “just in case”; optional improvements belong in plain-language prose only, "
    "without suggestion JSON, unless the analyst opted into alternates as above.\n"
)


def _parse_new_rows(payload: Dict[str, Any]) -> List[Dict[str, str]]:
    raw = payload.get("new_rows") or []
    if not isinstance(raw, list):
        return []
    out: List[Dict[str, str]] = []
    for item in raw[:8]:
        if not isinstance(item, dict):
            continue
        claim = str(item.get("claim") or "").strip()
        evidence = str(item.get("evidence") or "").strip()
        reasoning = str(item.get("reasoning") or "").strip()
        if not (claim or evidence or reasoning):
            continue
        out.append({"claim": claim, "evidence": evidence, "reasoning": reasoning})
    return out


def _normalize_suggestion_list(suggestions: Any) -> List[Dict[str, Any]]:
    if not isinstance(suggestions, list):
        return []
    filtered: List[Dict[str, Any]] = []
    for s in suggestions:
        if not isinstance(s, dict):
            continue
        field = s.get("field")
        if field not in ("claim", "evidence", "reasoning"):
            continue
        try:
            rid = int(s.get("row_id"))
        except Exception:
            continue
        filtered.append(
            {
                "row_id": rid,
                "field": field,
                "old_text": s.get("old_text") or "",
                "new_text": s.get("new_text") or "",
            }
        )
    return filtered


def _parse_suggestion_payload(payload: Any) -> Tuple[List[Dict[str, Any]], List[Dict[str, str]]]:
    if not isinstance(payload, dict):
        return [], []
    return _normalize_suggestion_list(payload.get("suggestions")), _parse_new_rows(payload)


def _json_candidate_strings(text: str) -> List[str]:
    """Collect JSON blobs that may contain lumenci suggestion schema (models often skip XML tags)."""
    text = text or ""
    out: List[str] = []
    lo = text.lower()
    otag = "<lumenci_suggestion_json>"
    ctag = "</lumenci_suggestion_json>"
    i = lo.find(otag)
    if i != -1:
        j = lo.find(ctag, i)
        if j != -1:
            inner = text[i + len(otag) : j].strip()
            if inner:
                out.append(inner)
    for m in re.finditer(r"```(?:json)?\s*([\s\S]*?)```", text, re.IGNORECASE):
        blob = m.group(1).strip()
        if blob and "suggestions" in blob:
            out.append(blob)
    for m in re.finditer(r"\{\s*\"suggestions\"\s*:", text):
        start = m.start()
        depth = 0
        for pos in range(start, len(text)):
            c = text[pos]
            if c == "{":
                depth += 1
            elif c == "}":
                depth -= 1
                if depth == 0:
                    out.append(text[start : pos + 1])
                    break
    seen = set()
    uniq: List[str] = []
    for b in out:
        if b in seen:
            continue
        seen.add(b)
        uniq.append(b)
    return uniq


def _strip_machine_json_for_display(text: str) -> str:
    """Keep chat readable: drop tagged blocks and obvious JSON code fences."""
    t = text or ""
    lo = t.lower()
    otag = "<lumenci_suggestion_json>"
    ctag = "</lumenci_suggestion_json>"
    i = lo.find(otag)
    if i != -1:
        j = lo.find(ctag, i)
        if j != -1:
            t = (t[:i] + t[j + len(ctag) :]).strip()
    t = re.sub(r"```(?:json)?\s*[\s\S]*?```", "", t, flags=re.IGNORECASE)
    return (t or "").strip()


def _extract_lumenci_payload(text: str) -> Tuple[str, List[Dict], List[Dict]]:
    raw = text or ""
    suggestions: List[Dict] = []
    new_rows: List[Dict[str, str]] = []
    for blob in _json_candidate_strings(raw):
        try:
            payload = json.loads(blob)
        except Exception:
            continue
        sug, nr = _parse_suggestion_payload(payload)
        if sug or nr:
            suggestions, new_rows = sug, nr
            break
    display = _strip_machine_json_for_display(raw)
    if not display:
        display = raw.strip()
    return display, suggestions, new_rows


def _recover_suggestions_from_prior_assistant(
    ch: ClaimChart, user_message: str, suggestions: List[Dict[str, Any]], new_rows: List[Dict[str, str]]
) -> Tuple[List[Dict[str, Any]], List[Dict[str, str]]]:
    """
    When the analyst says \"do it\" / \"apply\" the model sometimes replies with naked JSON or empty
    text. Re-use structured suggestions from the latest assistant message in the thread.
    """
    if suggestions:
        return suggestions, new_rows
    if not _user_confirms_apply(user_message):
        return suggestions, new_rows
    prev = (
        ChatMessage.objects.filter(claim_chart=ch, role=ChatMessage.Role.ASSISTANT)
        .order_by("-created_at", "-id")
        .first()
    )
    if not prev or not (prev.content or "").strip():
        return suggestions, new_rows
    _, s2, n2 = _extract_lumenci_payload(prev.content)
    return (s2 or suggestions), (n2 or new_rows)


def _user_wants_alternate_or_broad_edit(msg: str) -> bool:
    """True if analyst explicitly invites edits on already-strong rows or whole-chart rework."""
    ml = (msg or "").strip().lower()
    if not ml:
        return False
    phrases = (
        "alternate",
        "alternative",
        "another angle",
        "parallel argument",
        "parallel ",
        "different framing",
        "reframe",
        "optional wording",
        "optional reword",
        "optional alternative",
        "second option",
        "other way to say",
        "other way ",
        "all rows",
        "every row",
        "whole chart",
        "entire chart",
        "review all",
        "rewrite all",
        "revise all",
        "go through each",
        "same strength",
        "same level",
        "equally strong",
        "optional improvement",
        "brainstorm another",
    )
    return any(p in ml for p in phrases)


def _user_confirms_apply(msg: str) -> bool:
    """Short imperative messages that mean 'apply what you proposed' (bypass strength gate for those edits)."""
    ml = (msg or "").strip().lower()
    if not ml or len(ml) > 200:
        return False
    triggers = (
        "do it",
        "please do",
        "do the change",
        "do the changes",
        "make the change",
        "make those changes",
        "apply",
        "apply it",
        "apply that",
        "go ahead",
        "yes please",
        "implement",
        "update the table",
        "update the chart",
        "put it in the table",
        "save it",
        "yes do",
        "do that",
        "proceed",
        "commit",
    )
    return any(t in ml for t in triggers)


def _filter_suggestions_by_strength(
    suggestions: List[Dict[str, Any]],
    rows_by_id: Dict[int, ClaimChartRow],
    user_message: str,
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    """
    Drop structured edits targeting rows the analyst marked Strong, unless they asked for alternates /
    broad review. Acts as a safety net on top of the model instructions.
    """
    if not suggestions:
        return [], []
    if _user_wants_alternate_or_broad_edit(user_message) or _user_confirms_apply(user_message):
        return list(suggestions), []
    kept: List[Dict[str, Any]] = []
    removed: List[Dict[str, Any]] = []
    for s in suggestions:
        try:
            rid = int(s.get("row_id"))
        except (TypeError, ValueError):
            continue
        row = rows_by_id.get(rid)
        if not row:
            continue
        if row.strength == ClaimChartRow.Strength.STRONG:
            removed.append({"row_id": rid, "field": s.get("field")})
            continue
        kept.append(s)
    return kept, removed


def _history_controls(ch: ClaimChart) -> Dict[str, bool]:
    return {
        "can_undo": RowChange.objects.filter(claim_chart=ch, is_undone=False).exists(),
        "can_redo": RowChange.objects.filter(
            claim_chart=ch, is_undone=True, redo_invalidated=False
        ).exists(),
    }


def _invalidate_redo_branch(ch: ClaimChart) -> None:
    RowChange.objects.filter(claim_chart=ch, is_undone=True).update(redo_invalidated=True)


def _row_change_to_history_item(c: RowChange) -> Dict:
    if c.is_undone:
        status = "superseded" if c.redo_invalidated else "undone"
    else:
        status = "applied"
    field = c.field
    if field == "add_row":
        summary = f"Added row {c.row_index} (new claim element)"
        preview = ""
        try:
            payload = json.loads(c.new_text or "{}")
            preview = str(payload.get("claim") or "").strip().replace("\n", " ")
        except Exception:
            preview = ""
    else:
        summary = f"Row {c.row_index}: updated {field.replace('_', ' ')}"
        preview = (c.new_text or "").strip().replace("\n", " ")
    if len(preview) > 160:
        preview = preview[:157] + "…"
    return {
        "id": c.id,
        "row_index": c.row_index,
        "field": c.field,
        "status": status,
        "summary": summary,
        "preview": preview,
        "created_at": c.created_at.isoformat(),
        "undone_at": c.undone_at.isoformat() if c.undone_at else None,
    }


def _chart_to_dict(chart: ClaimChart) -> Dict:
    rows = []
    for r in chart.rows.all().order_by("row_index"):
        rows.append(
            {
                "id": r.row_index,
                "origin": r.origin,
                "strength": r.strength,
                "claim": r.claim_text,
                "evidence": r.evidence_text,
                "reasoning": r.reasoning_text,
            }
        )
    chat = []
    for m in chart.chat_messages.all().order_by("created_at", "id"):
        chat.append({"role": m.role, "content": m.content, "ts": int(m.created_at.timestamp() * 1000)})
    return {
        "id": chart.id,
        "case_id": chart.case_id,
        "name": chart.name,
        "source_type": chart.source_type,
        "status": chart.status,
        "error_message": chart.error_message,
        "system_instructions": chart.system_instructions,
        "source_file_url": _safe_file_url(chart.source_file),
        "rows": rows,
        "chat": chat,
        "history_controls": _history_controls(chart),
    }


@require_GET
def index(request):
    # UI now loads dynamically from APIs (no seeded state).
    return render(request, "assistant/index_db.html", {})


def _safe_file_url(f) -> str:
    if not f:
        return ""
    try:
        return f.url
    except ValueError:
        return ""


def _product_doc_brief(d: ProductDoc) -> Dict[str, Any]:
    return {
        "id": d.id,
        "name": d.name,
        "doc_type": d.doc_type,
        "claim_chart_id": d.claim_chart_id,
        "created_at": d.created_at.isoformat(),
        "file_url": _safe_file_url(d.file),
        "source_url": (d.source_url or "").strip(),
    }


@require_GET
def api_cases(request):
    cases = []
    qs = (
        Case.objects.all()
        .order_by("-created_at")
        .prefetch_related(
            Prefetch(
                "claim_charts",
                queryset=ClaimChart.objects.order_by("-created_at").prefetch_related(
                    Prefetch(
                        "product_docs",
                        queryset=ProductDoc.objects.order_by("-created_at"),
                    )
                ),
            )
        )
    )
    for c in qs:
        charts_out: List[Dict[str, Any]] = []
        for ch in c.claim_charts.all():
            charts_out.append(
                {
                    "id": ch.id,
                    "name": ch.name,
                    "status": ch.status,
                    "source_type": ch.source_type,
                    "created_at": ch.created_at.isoformat(),
                    "source_file_url": _safe_file_url(ch.source_file),
                    "product_docs": [_product_doc_brief(d) for d in ch.product_docs.all()],
                }
            )
        loose = c.product_docs.filter(claim_chart__isnull=True).order_by("-created_at")
        cases.append(
            {
                "id": c.id,
                "name": c.name,
                "created_at": c.created_at.isoformat(),
                "claim_charts": charts_out,
                "unassigned_product_docs": [_product_doc_brief(d) for d in loose],
                "product_docs": [_product_doc_brief(d) for d in c.product_docs.all().order_by("-created_at")],
            }
        )
    return JsonResponse({"ok": True, "cases": cases})


@require_POST
@csrf_exempt
def api_cases_create(request):
    try:
        body = json.loads(request.body.decode("utf-8") or "{}")
    except Exception:
        body = {}
    name = (body.get("name") or "").strip() or "New Case"
    c = Case.objects.create(name=name)
    return JsonResponse({"ok": True, "case": {"id": c.id, "name": c.name}})


@require_POST
@csrf_exempt
def api_case_update(request, case_id: int):
    try:
        body = json.loads(request.body.decode("utf-8") or "{}")
    except Exception:
        body = {}
    name = (body.get("name") or "").strip()
    case = get_object_or_404(Case, id=case_id)
    if name:
        case.name = name
        case.save(update_fields=["name"])
    return JsonResponse({"ok": True, "case": {"id": case.id, "name": case.name}})


@require_POST
@csrf_exempt
def api_case_delete(request, case_id: int):
    case = get_object_or_404(Case, id=case_id)
    case.delete()
    return JsonResponse({"ok": True})


@require_POST
@csrf_exempt
def api_claim_chart_update(request, chart_id: int):
    try:
        body = json.loads(request.body.decode("utf-8") or "{}")
    except Exception:
        body = {}
    name = (body.get("name") or "").strip()
    ch = get_object_or_404(ClaimChart, id=chart_id)
    fields: List[str] = []
    if name:
        ch.name = name
        fields.append("name")
    if "systemInstructions" in body:
        ch.system_instructions = (body.get("systemInstructions") or "").strip()
        fields.append("system_instructions")
    if fields:
        ch.save(update_fields=fields)
    return JsonResponse({"ok": True, "claim_chart": _chart_to_dict(ch)})


@require_POST
@csrf_exempt
def api_product_doc_update(request, doc_id: int):
    try:
        body = json.loads(request.body.decode("utf-8") or "{}")
    except Exception:
        body = {}
    name = (body.get("name") or "").strip()
    d = get_object_or_404(ProductDoc, id=doc_id)
    fields: List[str] = []
    if name:
        d.name = name
        fields.append("name")
    if "claim_chart_id" in body:
        raw = body.get("claim_chart_id")
        if raw in (None, "", 0, "0"):
            d.claim_chart = None
        else:
            ch = get_object_or_404(ClaimChart, id=int(raw), case_id=d.case_id)
            d.claim_chart = ch
        fields.append("claim_chart")
    if fields:
        d.save(update_fields=fields)
    return JsonResponse(
        {
            "ok": True,
            "product_doc": {
                "id": d.id,
                "name": d.name,
                "claim_chart_id": d.claim_chart_id,
            },
        }
    )


@require_POST
@csrf_exempt
def api_claim_chart_row_update(request, chart_id: int):
    ch = get_object_or_404(ClaimChart, id=chart_id)
    try:
        body = json.loads(request.body.decode("utf-8") or "{}")
    except Exception:
        body = {}
    try:
        row_index = int(body.get("row_index"))
    except Exception:
        return JsonResponse({"ok": False, "error": "row_index required"}, status=400)
    row = get_object_or_404(ClaimChartRow, claim_chart=ch, row_index=row_index)
    upd = []
    if "claim" in body:
        row.claim_text = str(body.get("claim") or "")
        upd.append("claim_text")
    if "evidence" in body:
        row.evidence_text = str(body.get("evidence") or "")
        upd.append("evidence_text")
    if "reasoning" in body:
        row.reasoning_text = str(body.get("reasoning") or "")
        upd.append("reasoning_text")
    st = body.get("strength")
    if st in ("strong", "weak", "missing"):
        row.strength = st
        upd.append("strength")
    og = body.get("origin")
    if og in (ClaimChartRow.RowOrigin.UPLOAD, ClaimChartRow.RowOrigin.ADDED):
        row.origin = og
        upd.append("origin")
    if upd:
        row.save(update_fields=upd)
    run_llm_strength = st not in ("strong", "weak", "missing")
    if run_llm_strength:
        sync_one_row_strength(row)
    ch = ClaimChart.objects.prefetch_related("rows", "chat_messages").get(pk=ch.pk)
    return JsonResponse({"ok": True, "claim_chart": _chart_to_dict(ch)})


@require_POST
@csrf_exempt
def api_claim_chart_row_delete(request, chart_id: int):
    ch = get_object_or_404(ClaimChart, id=chart_id)
    try:
        body = json.loads(request.body.decode("utf-8") or "{}")
    except Exception:
        body = {}
    try:
        row_index = int(body.get("row_index"))
    except Exception:
        return JsonResponse({"ok": False, "error": "row_index required"}, status=400)
    with transaction.atomic():
        ClaimChartRow.objects.filter(claim_chart=ch, row_index=row_index).delete()
        RowChange.objects.filter(claim_chart=ch, row_index=row_index).delete()
    ch = ClaimChart.objects.prefetch_related("rows", "chat_messages").get(pk=ch.pk)
    return JsonResponse({"ok": True, "claim_chart": _chart_to_dict(ch)})


@require_POST
@csrf_exempt
def api_claim_chart_row_add_empty(request, chart_id: int):
    ch = get_object_or_404(ClaimChart, id=chart_id)
    next_idx = (ch.rows.aggregate(m=Max("row_index"))["m"] or 0) + 1
    ClaimChartRow.objects.create(
        claim_chart=ch,
        row_index=next_idx,
        origin=ClaimChartRow.RowOrigin.ADDED,
        claim_text="",
        evidence_text="",
        reasoning_text="",
    )
    ch = ClaimChart.objects.prefetch_related("rows", "chat_messages").get(pk=ch.pk)
    return JsonResponse({"ok": True, "claim_chart": _chart_to_dict(ch)})


@require_POST
@csrf_exempt
def api_claim_chart_chat_clear(request, chart_id: int):
    ch = get_object_or_404(ClaimChart, id=chart_id)
    ChatMessage.objects.filter(claim_chart=ch).delete()
    ch = ClaimChart.objects.prefetch_related("rows", "chat_messages").get(pk=ch.pk)
    return JsonResponse({"ok": True, "claim_chart": _chart_to_dict(ch)})


@require_POST
@csrf_exempt
def api_claim_chart_history_clear(request, chart_id: int):
    ch = get_object_or_404(ClaimChart, id=chart_id)
    RowChange.objects.filter(claim_chart=ch).delete()
    ch = ClaimChart.objects.prefetch_related("rows", "chat_messages").get(pk=ch.pk)
    return JsonResponse({"ok": True, "claim_chart": _chart_to_dict(ch)})


@require_POST
@csrf_exempt
def api_claim_charts_upload(request):
    case_id = request.POST.get("case_id")
    f = request.FILES.get("file")
    if not case_id or not f:
        msg = "Missing case_id or file. Choose a matter first, then upload again."
        if settings.DEBUG:
            msg += f" (POST keys: {list(request.POST.keys())}, FILES: {list(request.FILES.keys())})"
        return JsonResponse({"ok": False, "error": msg}, status=400)

    case = get_object_or_404(Case, id=int(case_id))
    name = request.POST.get("name") or f.name
    ext = (os.path.splitext(f.name)[1] or "").lstrip(".").lower()
    ch = ClaimChart.objects.create(
        case=case,
        name=name,
        source_file=f,
        source_type=ext,
        status=ClaimChart.Status.UPLOADED,
    )
    # Parse immediately for demo (synchronous).
    try:
        ch.status = ClaimChart.Status.PARSING
        ch.error_message = ""
        ch.save(update_fields=["status", "error_message"])

        parsed = parse_claim_chart(ch.source_file.path)
        ClaimChartRow.objects.filter(claim_chart=ch).delete()
        for pr in parsed:
            ClaimChartRow.objects.create(
                claim_chart=ch,
                row_index=pr.row_index,
                origin=ClaimChartRow.RowOrigin.UPLOAD,
                claim_text=pr.claim_text,
                evidence_text=pr.evidence_text,
                reasoning_text=pr.reasoning_text,
            )

        sync_claim_chart_strengths(ch)
        ch.status = ClaimChart.Status.READY
        ch.save(update_fields=["status"])
    except ParseError as e:
        ch.status = ClaimChart.Status.ERROR
        ch.error_message = str(e)
        ch.save(update_fields=["status", "error_message"])
    except Exception as e:  # pragma: no cover
        ch.status = ClaimChart.Status.ERROR
        ch.error_message = f"Unexpected parse error: {type(e).__name__}"
        ch.save(update_fields=["status", "error_message"])
    return JsonResponse({"ok": True, "claim_chart": {"id": ch.id, "name": ch.name, "status": ch.status}})


@require_POST
@csrf_exempt
def api_product_docs_upload(request):
    case_id = request.POST.get("case_id")
    f = request.FILES.get("file")
    if not case_id or not f:
        msg = "Missing case_id or file. Choose a matter and chart first."
        if settings.DEBUG:
            msg += f" (POST keys: {list(request.POST.keys())}, FILES: {list(request.FILES.keys())})"
        return JsonResponse({"ok": False, "error": msg}, status=400)
    case = get_object_or_404(Case, id=int(case_id))
    chart_ref: Optional[ClaimChart] = None
    chart_id_raw = str(request.POST.get("claim_chart_id") or "").strip()
    if chart_id_raw:
        chart_ref = get_object_or_404(ClaimChart, id=int(chart_id_raw), case_id=case.id)
    else:
        # If the client omits chart id (older UI / gesture timing), bind to the only chart in the case when unambiguous.
        only = list(case.claim_charts.order_by("-created_at")[:2])
        if len(only) == 1:
            chart_ref = only[0]
    name = request.POST.get("name") or f.name
    ext = (os.path.splitext(f.name)[1] or "").lstrip(".").lower()
    d = ProductDoc.objects.create(case=case, claim_chart=chart_ref, name=name, file=f, doc_type=ext)
    # Best-effort extraction for chat context (RAG-lite).
    try:
        text = extract_product_doc_text(d.file.path)
        d.extracted_text = (text or "").strip()
        d.extracted_error = ""
        d.extracted_at = timezone.now()
        d.save(update_fields=["extracted_text", "extracted_error", "extracted_at"])
    except Exception as e:
        d.extracted_text = ""
        d.extracted_error = f"{type(e).__name__}: {str(e)}"
        d.extracted_at = timezone.now()
        d.save(update_fields=["extracted_text", "extracted_error", "extracted_at"])
    return JsonResponse(
        {
            "ok": True,
            "product_doc": {"id": d.id, "name": d.name, "claim_chart_id": d.claim_chart_id},
        }
    )


@require_POST
@csrf_exempt
def api_product_docs_from_url(request):
    try:
        body = json.loads(request.body.decode("utf-8") or "{}")
    except Exception:
        body = {}
    case_id = body.get("case_id")
    url_raw = str(body.get("url") or "").strip()
    if case_id in (None, "", 0, "0") or not url_raw:
        return JsonResponse({"ok": False, "error": "case_id and url required"}, status=400)
    case = get_object_or_404(Case, id=int(case_id))
    chart_ref: Optional[ClaimChart] = None
    chart_id_raw = str(body.get("claim_chart_id") or "").strip()
    if chart_id_raw:
        chart_ref = get_object_or_404(ClaimChart, id=int(chart_id_raw), case_id=case.id)
    else:
        only = list(case.claim_charts.order_by("-created_at")[:2])
        if len(only) == 1:
            chart_ref = only[0]

    from .url_scrape import URLFetchError, fetch_page_text

    if not chart_ref:
        return JsonResponse(
            {
                "ok": False,
                "error": "Select which claim chart this URL supports (active chart ⇄, or only one chart in the matter).",
            },
            status=400,
        )

    try:
        final_url, title, text = fetch_page_text(url_raw)
    except URLFetchError as e:
        return JsonResponse({"ok": False, "error": str(e)}, status=400)
    except Exception as e:  # pragma: no cover
        return JsonResponse({"ok": False, "error": f"{type(e).__name__}: {e}"}, status=500)

    display = (title or "").strip()[:180]
    if not display:
        from urllib.parse import urlparse

        p = urlparse(final_url)
        display = ((p.netloc or "") + (p.path or ""))[:180] or "Web page"
    name = f"Web: {display}"

    extracted = (text or "").strip()
    extracted_err = ""
    if not extracted:
        extracted_err = "No readable text could be extracted from this page."

    d = ProductDoc.objects.create(
        case=case,
        claim_chart=chart_ref,
        name=name,
        file=None,
        doc_type="url",
        source_url=final_url[:2048],
        extracted_text=extracted,
        extracted_error=extracted_err,
        extracted_at=timezone.now(),
    )
    return JsonResponse(
        {
            "ok": True,
            "product_doc": {
                "id": d.id,
                "name": d.name,
                "claim_chart_id": d.claim_chart_id,
                "source_url": d.source_url,
            },
        }
    )


def _chunk_text(text: str, chunk_size: int = 1200, overlap: int = 180) -> List[str]:
    t = (text or "").strip()
    if not t:
        return []
    out: List[str] = []
    i = 0
    n = len(t)
    while i < n and len(out) < 24:
        j = min(n, i + chunk_size)
        out.append(t[i:j])
        if j >= n:
            break
        i = max(0, j - overlap)
    return out


def _score_chunk(chunk: str, needles: List[str]) -> int:
    c = (chunk or "").lower()
    score = 0
    for w in needles:
        if not w or len(w) < 4:
            continue
        if w in c:
            score += 3
    return score


def _build_doc_context(docs: List[ProductDoc], user_message: str, chart_rows: List[Dict[str, Any]]) -> str:
    # Rank doc chunks by simple keyword overlap (fast + zero extra deps).
    needles: List[str] = []
    um = (user_message or "").lower()
    for tok in re.split(r"[^a-z0-9_]+", um):
        if len(tok) >= 4:
            needles.append(tok)
    # Add some claim terms too (helps when the user says "fix weak reasoning" without specifics)
    for r in chart_rows[:18]:
        for field in ("claim", "evidence", "reasoning"):
            v = str(r.get(field) or "").lower()
            for tok in re.split(r"[^a-z0-9_]+", v):
                if 6 <= len(tok) <= 24:
                    needles.append(tok)
    needles = list(dict.fromkeys(needles))[:60]

    scored: List[Tuple[int, str]] = []
    for d in docs:
        if not (d.extracted_text or "").strip():
            continue
        for ch in _chunk_text(d.extracted_text):
            s = _score_chunk(ch, needles)
            if s <= 0:
                continue
            scored.append((s, f"[{d.name}]\n{ch}"))
    scored.sort(key=lambda x: x[0], reverse=True)
    top = [t for _, t in scored[:4]]
    if not top:
        available = [d.name for d in docs if (d.extracted_text or "").strip()]
        if available:
            return "Product doc text is available, but no highly relevant snippets matched the latest user message.\n"
        return ""
    return "\n\n---\n\n".join(top)


@require_POST
@csrf_exempt
def api_product_doc_delete(request, doc_id: int):
    d = get_object_or_404(ProductDoc, id=doc_id)
    try:
        if d.file:
            d.file.delete(save=False)
    except Exception:
        pass
    d.delete()
    return JsonResponse({"ok": True})


@require_POST
@csrf_exempt
def api_claim_chart_delete(request, chart_id: int):
    ch = get_object_or_404(ClaimChart, id=chart_id)
    try:
        if ch.source_file:
            ch.source_file.delete(save=False)
    except Exception:
        pass
    ch.delete()
    return JsonResponse({"ok": True})


@require_GET
def api_claim_chart_detail(request, chart_id: int):
    ch = get_object_or_404(ClaimChart.objects.prefetch_related("rows", "chat_messages"), id=chart_id)
    return JsonResponse({"ok": True, "claim_chart": _chart_to_dict(ch)})


def _safe_docx_filename(name: str) -> str:
    base = (name or "claim-chart").strip() or "claim-chart"
    for c in '<>:"/\\|?*\n\r':
        base = base.replace(c, "_")
    base = base[:150]
    if not base.lower().endswith(".docx"):
        base = f"{base}.docx"
    return base


@require_GET
def api_claim_chart_export_docx(request, chart_id: int):
    ch = get_object_or_404(ClaimChart.objects.select_related("case"), id=chart_id)
    # Always read rows fresh from DB (no stale prefetch); order matches on-screen chart.
    rows_qs = ClaimChartRow.objects.filter(claim_chart_id=ch.pk).order_by("row_index")
    doc = Document()
    doc.add_heading(ch.name or "Claim chart", level=1)
    meta = doc.add_paragraph()
    meta.add_run("Case: ").bold = True
    meta.add_run(ch.case.name if ch.case else "—")
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "#"
    hdr_cells[1].text = "Claim limitation / element"
    hdr_cells[2].text = "Evidence (prior art / exhibit)"
    hdr_cells[3].text = "How the reference teaches / discloses"
    for row in rows_qs:
        cells = table.add_row().cells
        cells[0].text = str(row.row_index)
        cells[1].text = row.claim_text or ""
        cells[2].text = row.evidence_text or ""
        cells[3].text = row.reasoning_text or ""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = _safe_docx_filename(ch.name or "claim-chart")
    resp = FileResponse(
        buf,
        as_attachment=True,
        filename=fname,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    # Browsers often cache GET downloads; avoid serving an old .docx after the chart changed.
    resp["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp["Pragma"] = "no-cache"
    return resp


@require_POST
@csrf_exempt
def api_claim_chart_apply_suggestion(request, chart_id: int):
    ch = get_object_or_404(ClaimChart.objects.select_related("case"), id=chart_id)
    try:
        body = json.loads(request.body.decode("utf-8") or "{}")
    except Exception:
        body = {}
    suggestion = body.get("suggestion") or {}
    if (suggestion.get("type") or "").strip().lower() == "add_row":
        claim = (suggestion.get("claim") or "").strip()
        evidence = (suggestion.get("evidence") or "").strip()
        reasoning = (suggestion.get("reasoning") or "").strip()
        if not (claim or evidence or reasoning):
            return JsonResponse({"ok": False, "error": "New row must have at least one column"}, status=400)
        _invalidate_redo_branch(ch)
        next_idx = (ch.rows.aggregate(m=Max("row_index"))["m"] or 0) + 1
        snapshot = json.dumps(
            {
                "claim": claim,
                "evidence": evidence,
                "reasoning": reasoning,
                "origin": ClaimChartRow.RowOrigin.ADDED,
            },
            ensure_ascii=False,
        )
        with transaction.atomic():
            RowChange.objects.create(
                claim_chart=ch,
                row_index=next_idx,
                field="add_row",
                old_text="",
                new_text=snapshot,
            )
            new_row = ClaimChartRow.objects.create(
                claim_chart=ch,
                row_index=next_idx,
                origin=ClaimChartRow.RowOrigin.ADDED,
                claim_text=claim,
                evidence_text=evidence,
                reasoning_text=reasoning,
            )
        sync_one_row_strength(new_row)
        ch = ClaimChart.objects.prefetch_related("rows", "chat_messages").get(pk=ch.pk)
        return JsonResponse({"ok": True, "claim_chart": _chart_to_dict(ch)})

    row_id = int(suggestion.get("row_id") or 0)
    field = suggestion.get("field")
    new_text = suggestion.get("new_text") or ""
    if row_id <= 0 or field not in ("claim", "evidence", "reasoning"):
        return JsonResponse({"ok": False, "error": "Invalid suggestion"}, status=400)

    _invalidate_redo_branch(ch)
    row = get_object_or_404(ClaimChartRow, claim_chart=ch, row_index=row_id)
    old_text = (
        row.claim_text
        if field == "claim"
        else row.evidence_text
        if field == "evidence"
        else row.reasoning_text
    )
    with transaction.atomic():
        RowChange.objects.create(
            claim_chart=ch, row_index=row_id, field=field, old_text=old_text, new_text=new_text
        )
        if field == "claim":
            row.claim_text = new_text
        elif field == "evidence":
            row.evidence_text = new_text
        else:
            row.reasoning_text = new_text
        row.save(update_fields=["claim_text", "evidence_text", "reasoning_text"])

    sync_one_row_strength(row)
    ch = ClaimChart.objects.prefetch_related("rows", "chat_messages").get(pk=ch.pk)
    return JsonResponse({"ok": True, "claim_chart": _chart_to_dict(ch)})


@require_POST
@csrf_exempt
def api_claim_chart_undo(request, chart_id: int):
    ch = get_object_or_404(ClaimChart, id=chart_id)
    last = (
        RowChange.objects.filter(claim_chart=ch, is_undone=False)
        .order_by("-created_at", "-id")
        .first()
    )
    if not last:
        return JsonResponse({"ok": True, "message": "Nothing to undo", "claim_chart": _chart_to_dict(ch)})

    row_to_resync = None
    with transaction.atomic():
        if last.field == "add_row":
            ClaimChartRow.objects.filter(claim_chart=ch, row_index=last.row_index).delete()
        else:
            row = get_object_or_404(ClaimChartRow, claim_chart=ch, row_index=last.row_index)
            if last.field == "claim":
                row.claim_text = last.old_text
            elif last.field == "evidence":
                row.evidence_text = last.old_text
            else:
                row.reasoning_text = last.old_text
            row.save(update_fields=["claim_text", "evidence_text", "reasoning_text"])
            row_to_resync = row
        last.is_undone = True
        last.undone_at = timezone.now()
        last.save(update_fields=["is_undone", "undone_at"])

    if row_to_resync is not None:
        sync_one_row_strength(row_to_resync)
    ch = ClaimChart.objects.prefetch_related("rows", "chat_messages").get(pk=ch.pk)
    return JsonResponse({"ok": True, "message": "Reverted", "claim_chart": _chart_to_dict(ch)})


@require_POST
@csrf_exempt
def api_claim_chart_redo(request, chart_id: int):
    ch = get_object_or_404(ClaimChart, id=chart_id)
    redo_op = (
        RowChange.objects.filter(claim_chart=ch, is_undone=True, redo_invalidated=False)
        .order_by("-undone_at", "-id")
        .first()
    )
    if not redo_op:
        return JsonResponse({"ok": True, "message": "Nothing to redo", "claim_chart": _chart_to_dict(ch)})

    row_to_resync = None
    with transaction.atomic():
        if redo_op.field == "add_row":
            try:
                payload = json.loads(redo_op.new_text or "{}")
            except json.JSONDecodeError:
                payload = {}
            claim = str(payload.get("claim") or "")
            evidence = str(payload.get("evidence") or "")
            reasoning = str(payload.get("reasoning") or "")
            origin = payload.get("origin") or ClaimChartRow.RowOrigin.ADDED
            if origin not in (ClaimChartRow.RowOrigin.UPLOAD, ClaimChartRow.RowOrigin.ADDED):
                origin = ClaimChartRow.RowOrigin.ADDED
            new_row = ClaimChartRow.objects.create(
                claim_chart=ch,
                row_index=redo_op.row_index,
                origin=origin,
                claim_text=claim,
                evidence_text=evidence,
                reasoning_text=reasoning,
            )
            row_to_resync = new_row
        else:
            row = get_object_or_404(ClaimChartRow, claim_chart=ch, row_index=redo_op.row_index)
            if redo_op.field == "claim":
                row.claim_text = redo_op.new_text
            elif redo_op.field == "evidence":
                row.evidence_text = redo_op.new_text
            else:
                row.reasoning_text = redo_op.new_text
            row.save(update_fields=["claim_text", "evidence_text", "reasoning_text"])
            row_to_resync = row
        redo_op.is_undone = False
        redo_op.undone_at = None
        redo_op.save(update_fields=["is_undone", "undone_at"])

    if row_to_resync is not None:
        sync_one_row_strength(row_to_resync)
    ch = ClaimChart.objects.prefetch_related("rows", "chat_messages").get(pk=ch.pk)
    return JsonResponse({"ok": True, "message": "Redone", "claim_chart": _chart_to_dict(ch)})


@require_GET
def api_claim_chart_history(request, chart_id: int):
    ch = get_object_or_404(ClaimChart, id=chart_id)
    try:
        limit = max(1, min(int(request.GET.get("limit", 80)), 200))
    except Exception:
        limit = 80
    qs = RowChange.objects.filter(claim_chart=ch).order_by("-created_at", "-id")[:limit]
    items = [_row_change_to_history_item(c) for c in qs]
    return JsonResponse({"ok": True, "changes": items, "history_controls": _history_controls(ch)})


@require_POST
@csrf_exempt
def api_claim_chart_chat(request, chart_id: int):
    ch = get_object_or_404(
        ClaimChart.objects.select_related("case").prefetch_related(
            "rows", "chat_messages", "product_docs", "case__product_docs"
        ),
        id=chart_id,
    )
    try:
        body = json.loads(request.body.decode("utf-8") or "{}")
    except Exception:
        body = {}

    user_message = (body.get("message") or "").strip()
    system_instructions = (body.get("systemInstructions") or "").strip()
    if system_instructions:
        ch.system_instructions = system_instructions
        ch.save(update_fields=["system_instructions"])

    if user_message:
        ChatMessage.objects.create(claim_chart=ch, role=ChatMessage.Role.USER, content=user_message)

    api_key = getattr(settings, "GROQ_API_KEY", "") or ""
    if not api_key or Groq is None:
        assistant_content = (
            "Groq is not configured. Set `GROQ_API_KEY` in your environment and try again."
        )
        ChatMessage.objects.create(claim_chart=ch, role=ChatMessage.Role.ASSISTANT, content=assistant_content)
        return JsonResponse(
            {
                "ok": True,
                "assistant": assistant_content,
                "suggestions": [],
                "new_rows": [],
                "claim_chart": _chart_to_dict(ch),
                "strength_gate": {"removed": [], "removed_count": 0},
            }
        )

    model = getattr(settings, "GROQ_MODEL", None) or "llama-3.3-70b-versatile"
    client = Groq(api_key=api_key)

    chart_rows = [
        {
            "row_id": r.row_index,
            "strength": r.strength,
            "claim": r.claim_text,
            "evidence": r.evidence_text,
            "reasoning": r.reasoning_text,
        }
        for r in ch.rows.all().order_by("row_index")
    ]
    doc_qs = list(ch.product_docs.all().order_by("-created_at"))
    docs = [{"name": d.name, "doc_type": d.doc_type} for d in doc_qs]
    loose_names = list(
        ch.case.product_docs.filter(claim_chart__isnull=True)
        .order_by("-created_at")
        .values_list("name", flat=True)[:12]
    )

    extra_system = ""
    if ch.system_instructions:
        extra_system = f"\n\nAnalyst system instructions:\n{ch.system_instructions}\n"

    other_docs_note = ""
    if loose_names:
        other_docs_note = (
            "\n\nOther technical docs in this case are NOT linked to this claim chart (filenames only): "
            + json.dumps(loose_names, ensure_ascii=False)
            + "\nThey are excluded from extracted snippets unless the analyst links them to this chart in the UI.\n"
        )

    full_system = (
        SYSTEM_PROMPT
        + STRENGTH_POLICY
        + extra_system
        + "\nTechnical docs linked to THIS claim chart (filenames only):\n"
        + json.dumps(docs, ensure_ascii=False)
        + other_docs_note
        + "\n\nExtracted text snippets from docs linked to this chart (RAG-lite, may be partial):\n"
        + (_build_doc_context(doc_qs, user_message, chart_rows) or "—")
        + "\n\nCurrent claim chart JSON:\n"
        + json.dumps(chart_rows, ensure_ascii=False)
        + "\n\nReturn suggestions only for existing row_id values present in the chart JSON. "
        "Use new_rows only for genuinely new claim elements not represented above."
    )

    messages = []
    for m in ch.chat_messages.all().order_by("-created_at", "-id")[:20][::-1]:
        messages.append({"role": m.role, "content": m.content})

    assistant_content = ""
    suggestions: List[Dict[str, Any]] = []
    new_rows: List[Dict[str, str]] = []
    last_exc: Optional[BaseException] = None
    max_attempts = 3
    for attempt in range(max_attempts):
        try:
            resp = client.chat.completions.create(
                model=model,
                max_tokens=2048,
                temperature=0.2,
                messages=([{"role": "system", "content": full_system}] + messages),
            )
            assistant_raw = (
                (resp.choices[0].message.content if getattr(resp, "choices", None) else "") or ""
            ).strip()
            assistant_content, suggestions, new_rows = _extract_lumenci_payload(assistant_raw)
            suggestions, new_rows = _recover_suggestions_from_prior_assistant(
                ch, user_message, suggestions, new_rows
            )
            if not (assistant_content or "").strip() and suggestions and _user_confirms_apply(user_message):
                assistant_content = (
                    "Use **Accept** in the claim chart for each proposed change, or say what to adjust."
                )
            last_exc = None
            break
        except Exception as e:
            last_exc = e
            name = type(e).__name__
            is_rl = name == "RateLimitError"
            if is_rl and attempt < max_attempts - 1:
                time.sleep(_groq_retry_after_seconds(e))
                continue
            if is_rl:
                assistant_content = GROQ_RATE_LIMIT_USER_MESSAGE
            else:
                detail = str(e).strip()
                if detail and len(detail) < 300:
                    assistant_content = (
                        f"I hit an error while contacting Groq ({name}). {detail}"
                    )
                else:
                    assistant_content = f"I hit an error while contacting Groq. Error: {name}"
            suggestions = []
            new_rows = []
            break

    # Validate suggestions: must target existing rows and fields.
    valid_row_ids = {r.row_index for r in ch.rows.all()}
    validated = []
    rows_by_id = {r.row_index: r for r in ch.rows.all()}
    for s in suggestions:
        rid = int(s.get("row_id"))
        field = s.get("field")
        if rid not in valid_row_ids:
            continue
        if field not in ("claim", "evidence", "reasoning"):
            continue
        if not (s.get("old_text") or "").strip():
            row = rows_by_id.get(rid)
            if row:
                s["old_text"] = (
                    row.claim_text
                    if field == "claim"
                    else row.evidence_text
                    if field == "evidence"
                    else row.reasoning_text
                )
        validated.append(s)

    gated, strength_removed = _filter_suggestions_by_strength(validated, rows_by_id, user_message)

    ChatMessage.objects.create(claim_chart=ch, role=ChatMessage.Role.ASSISTANT, content=assistant_content)
    return JsonResponse(
        {
            "ok": True,
            "assistant": assistant_content,
            "suggestions": gated,
            "new_rows": new_rows,
            "claim_chart": _chart_to_dict(ch),
            "strength_gate": {
                "removed": strength_removed,
                "removed_count": len(strength_removed),
            },
        }
    )
