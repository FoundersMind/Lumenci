"""One-off script: generate MVP PRD as .docx for stakeholders."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "PRD_Lumenci_Spark_MVP.docx"


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h1 = doc.add_heading("Product Requirements Document: Lumenci Spark (MVP)", 0)
    h1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    r = meta.add_run("Version: 1.0  |  Owner: Product  |  Audience: Engineering, Design, Analyst stakeholders")
    r.italic = True

    doc.add_heading("1. Problem Statement", level=1)
    doc.add_paragraph(
        "Patent analysts maintain claim charts (claim element, accused-product evidence, reasoning) "
        "alongside technical product materials—often across spreadsheets, email, and generic chat tools. "
        'This creates friction: AI assistance lacks reliable access to the right technical documents, '
        'suggestions may overwrite analyst-approved "strong" analysis, and delivering a clean Word '
        'deliverable from the "current" chart is error-prone. Lumenci Spark unifies matter workspace, '
        "grounded chat refinement, and exportable outcomes in one analyst-focused workbench."
    )

    doc.add_heading("2. User Stories", level=1)
    stories = [
        "As a patent analyst, I want to upload a claim chart file and see a three-column grid so that I can "
        "review and edit claim elements, evidence, and reasoning in one place.",
        "As a patent analyst, I want to attach technical/product documents to a specific claim chart so that "
        "the copilot can ground answers in the correct product context (RAG-style).",
        "As a patent analyst, I want to type natural-language requests in chat (e.g., strengthen weak evidence, "
        "tighter reasoning) so that I receive concrete, row-level suggestions.",
        "As a patent analyst, I want to accept, reject, undo, or redo table changes so that only human-approved "
        "edits become the saved chart state.",
        'As a patent analyst, I want the system to avoid revising rows I marked "strong" unless I explicitly ask '
        "for alternate wording or a broader review so that good work is not churned unnecessarily.",
        "As a patent analyst, I want to export the saved chart to Word so that I can hand off a client-ready "
        "artifact that reflects the latest accepted content.",
    ]
    for s in stories:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("3. Core Features", level=1)
    doc.add_paragraph("In scope (MVP)", style="Heading 3")
    in_scope = [
        "Matter (case) workspace with multiple claim charts and chart selection.",
        "Claim chart upload + parsing; tabular UI with strength/origin metadata and row edit.",
        "Per-chart technical document uploads with server-side text extraction; document context passed into "
        "chat for the active chart only.",
        "Per-chart custom instructions (system-style guidance).",
        "Chat copilot with structured suggestions mapped to rows/fields; apply/reject in the grid.",
        "Undo/redo and edit history aligned to accepted changes.",
        "Export current saved chart to Word (.docx) with no stale/cached download.",
        "Resilience: rate-limit handling and clear errors for AI provider failures.",
    ]
    for x in in_scope:
        doc.add_paragraph(x, style="List Bullet")

    doc.add_paragraph("Out of scope (MVP)", style="Heading 3")
    out_scope = [
        "Enterprise IAM, multi-tenant admin, granular roles.",
        "Dedicated vector database, semantic search UI, or cross-matter knowledge graph.",
        "Automated legal conclusions (infringement/validity) beyond analyst-driven refinement.",
        "Native mobile apps; offline mode.",
        "Real-time multi-user co-editing.",
    ]
    for x in out_scope:
        doc.add_paragraph(x, style="List Bullet")

    doc.add_heading("4. Key Decisions", level=1)
    decisions = [
        (
            "Chart-scoped documents and context.",
            "Technical documents are linked to a claim chart (not only the matter) so chat RAG and UI grouping "
            "prevent the wrong exhibit from informing the wrong chart.",
        ),
        (
            "Strength gate with explicit override.",
            'Server-side filtering and prompt policy protect "strong" rows; users must signal appetite for '
            "alternates or full-chart review—reduces noise and trust erosion.",
        ),
        (
            "Saved database state is the export source of truth.",
            "Export reflects persisted rows after accepts; proposed/pending UI changes are excluded until "
            "confirmed—aligns export with analyst intent and auditability.",
        ),
    ]
    for title, body in decisions:
        p = doc.add_paragraph()
        p.add_run(title + " ").bold = True
        p.add_run(body)

    doc.add_heading("5. Acceptance Criteria", level=1)
    criteria = [
        "Given a valid chart file, the UI renders all rows across three columns with counts consistent with "
        "the source.",
        "Given a technical document linked to the active chart, a chat query can be answered using substance "
        "from that document—not filename-only placeholders.",
        "Given an AI suggestion for a row/field, the user can Accept or Reject; accepted updates persist after "
        "refresh; Undo reverts the last accepted change.",
        'Given rows marked strong, default refinement requests do not produce chart edits for those rows '
        "unless the user explicitly requests alternates or broad review.",
        "Given pending suggestions in the grid, Export Word contains only saved row content; the product "
        "communicates that pending changes must be accepted to be included.",
        "Given provider rate limits, the user sees a clear message and the system retries or backs off without "
        "silent failure where feasible.",
    ]
    for x in criteria:
        doc.add_paragraph(x, style="List Number")

    doc.add_heading("6. Success Metrics", level=1)
    metrics = [
        "Time to first accepted AI-assisted edit after matter/chart setup (median).",
        "Suggestion acceptance rate versus reject or dismiss (usefulness signal).",
        "Successful Word export sessions per active matter (delivery completion).",
        "Qualitative: analyst trust in doc-grounded answers; fewer issues with wrong-chart context or "
        'unwanted edits to "strong" rows.',
    ]
    for x in metrics:
        doc.add_paragraph(x, style="List Bullet")

    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run(
        "Note: Target one printed page; if pagination overflows, shorten examples in design reviews or reduce "
        "bullet granularity."
    )
    fr.italic = True

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
